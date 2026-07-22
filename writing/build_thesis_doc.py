"""Build a Word draft of thesis content from the work done so far.
Run with: uv run python build_thesis_doc.py
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

FIG_DIR = "/home/tobi/Documents/thesis/code/hpc-dl-enhanced/writing/figures"
OUT_PATH = "/home/tobi/Documents/thesis/code/hpc-dl-enhanced/writing/thesis_draft_notes.docx"

doc = Document()

# --- basic styling ---
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def p(text):
    doc.add_paragraph(text)

def fig(filename, caption):
    doc.add_picture(f"{FIG_DIR}/{filename}", width=Inches(5.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(10)

# ============================================================
doc.add_heading("Working Draft: Chapter 3 and Chapter 5 Material", level=0)
p("This document collects the writing-ready material produced while building the "
  "data pipeline for this thesis. It is organized by the chapter each section is "
  "expected to land in, following the structure recorded in WRITING_TRACKER.md. "
  "None of this is final prose — line numbers, wording, and even section order "
  "will change once the surrounding chapters exist — but the substance, the "
  "numbers, and the reasoning behind each decision are accurate as of this draft "
  "and worth keeping rather than reconstructing later from memory.")

# ============================================================
h1("Chapter 3: Datasets and Feature Engineering")

h2("3.1 Data Acquisition and Schema Verification")
p("Two datasets anchor this study: F-DATA, drawn from the Fugaku supercomputer, "
  "and PM100, drawn from CINECA's Marconi100 system. Both are published "
  "job-accounting traces, but they record different things, and I wanted that "
  "difference established from the schema itself rather than assumed from the "
  "papers describing them.")
p("F-DATA arrived as 38 monthly parquet files, one per month from March 2021 "
  "through April 2024, totaling 25,866,900 job records — a number I confirmed "
  "directly from parquet metadata rather than the approximately-24-million figure "
  "quoted in the dataset's own paper. The file count is worth stating carefully: "
  "when this pipeline first ran, only one of the 38 files had actually been "
  "downloaded, and the loading code printed how many of the expected 38 it found "
  "rather than assuming the full set was present. All 38 were confirmed present "
  "by listing the directory directly — March 2021 through April 2024 with no "
  "gaps — before any of the numbers in this chapter that depend on the full "
  "date range were computed. PM100 is a single file of 231,238 jobs from "
  "Marconi100, covering May through October 2020.")
p("The original project proposal described both systems as heterogeneous "
  "CPU-GPU workloads. That framing does not survive contact with the schema. "
  "Fugaku's A64FX processors are CPU-only — there is no GPU field to speak of "
  "in F-DATA — while PM100's nodes do carry GPUs. F-DATA's columns include six "
  "hardware performance counters (labeled perf1 through perf6), a FLOP count, "
  "memory bandwidth, and an operational-intensity field, none of which appear "
  "anywhere in PM100. I searched PM100's 35 columns for anything matching flop, "
  "perf, cycle, instr, mbwidth, or opint as a final check before ruling out "
  "Roofline-style analysis there, and found nothing. The two datasets are "
  "complementary, not symmetric: F-DATA supports genuine Roofline analysis "
  "because it has the hardware counters that method needs, and PM100 does not.")

h2("3.2 Target Definitions and a Units Discrepancy")
p("F-DATA's three prediction targets are duration (execution time), mmszu "
  "(memory used), and avgpcon (power). PM100's are run_time, mem_alloc, and "
  "node_power_consumption. PM100 has no field recording memory actually used, "
  "only requested and allocated amounts, so mem_alloc is not a preference "
  "between two options — it is the only one available.")
p("avgpcon needed a closer look before I trusted it as a target. Its documented "
  "name, average node power consumption, implies a single node's draw, but the "
  "maximum value in a 5,000-row sample was 259,097 watts — far beyond anything "
  "a single A64FX node could draw. Restricting to single-node jobs (nnuma == 1) "
  "brought the range down to 39.6–149.9 watts, consistent with published A64FX "
  "power figures, and avgpcon correlated with the number of allocated nodes at "
  "0.98. The field is a job-level total, summed across every allocated node and "
  "averaged over time, not a per-node figure. I checked the unit consistency a "
  "second way: dividing avgpcon multiplied by duration by the recorded energy "
  "field (econ) gave a ratio clustering around 1/3600 across the sample, which "
  "is exactly what you would expect if avgpcon is in watts, duration is in "
  "seconds, and econ is recorded in watt-hours rather than joules. The numbers "
  "are internally consistent once the units are understood correctly; the "
  "column name is just misleading about scope. This has one concrete "
  "consequence for later chapters: F-DATA's power target measures a whole "
  "job's total draw, while PM100's node_power_consumption is genuinely "
  "per-node. The two are not directly comparable, and any table that puts them "
  "side by side needs to say so.")

p("Execution time, memory, and power are all heavy-tailed in both datasets: "
  "most jobs are short and small, and a long tail of much larger jobs pulls "
  "the mean well above the median. F-DATA's duration in a 5,000-row sample "
  "ranges from 1 second to 255,677 seconds (about 71 hours), with a median "
  "around 2,000 seconds — the kind of spread that will dominate a "
  "mean-squared-error loss if targets are trained on in raw units. All six "
  "targets across both datasets are therefore modeled in log1p space, with "
  "metrics reported back in real units afterward for interpretability. I "
  "did not assume the transform was safe; expm1(log1p(x)) was checked "
  "against the original values for all six targets and matched to within "
  "floating-point precision.")
fig("nb01_cell15_img1.png",
    "Figure 3.4. F-DATA execution time, raw versus log1p-transformed, "
    "illustrating the heavy-tailed distribution that motivates training in "
    "log space.")
p("Only one month of F-DATA was available when this pipeline was first "
  "built; all 38 months were obtained shortly afterward. The two "
  "job-count-over-time plots below cover the full range now available for "
  "both datasets and will matter directly in Chapter 4, where the "
  "chronological train/test split depends on there being enough time span "
  "to split meaningfully.")
fig("nb01_cell18_img7.png", "Figure 3.5. F-DATA jobs per day across the sampled files.")
fig("nb01_cell18_img8.png", "Figure 3.6. PM100 jobs per day across its full six-month span.")

h2("3.3 Submission-Time and Post-Hoc Feature Tiers")
p("The proposal's original feature diagram feeds measured FLOPs, memory "
  "bandwidth, and power directly into a network that predicts duration and "
  "power for the same job. That is circular if the intended use case is "
  "predicting a job's behavior before it runs, since none of those measured "
  "quantities exist until the job has already executed.")
p("I split every feature into two tiers to make this distinction impossible to "
  "blur by accident. Tier A holds only what is known at submission time: "
  "requested cores, nodes, and memory; queue and priority; requested wall time; "
  "an anonymized job-name embedding; and, for F-DATA, a rolling average of each "
  "user's recent job durations computed using only jobs strictly before the "
  "current one. Tier B holds everything measured during or after execution — "
  "allocated versus used resources, exit codes, the hardware counters, and the "
  "power fields. Tier A is what a scheduler could actually use to predict a "
  "job's behavior before running it; Tier B exists for characterizing what "
  "happened after the fact and for the residual-learning step of the hybrid "
  "model described in Chapter 4. The two tiers are built by separate functions "
  "rather than one function with a flag, specifically so that mixing them "
  "requires deliberately writing code to do it. An assertion checks for column "
  "overlap between the two lists every time the module loads, and a second "
  "assertion re-checks after every transformation step in the feature pipeline, "
  "not just once at the start.")

h2("3.4 Embedding Dimensionality Reduction")
p("F-DATA's job-name field is anonymized and replaced with a 384-dimensional "
  "Sentence-BERT embedding. Feeding 384 raw dimensions into a random forest or "
  "gradient-boosted tree model is impractical, so the plan called for reducing "
  "it with PCA to, in the proposal's words, \"a handful of components.\" I had "
  "originally set that handful to 10 without checking what fraction of the "
  "embedding's variance 10 components actually capture.")
p("It turned out to be 68 percent on a single month's data — nowhere near "
  "enough to justify the number. I replaced the fixed default with an "
  "empirical check: fit PCA with up to 100 components, look at cumulative "
  "explained variance, and keep the smallest number that reaches a 90 percent "
  "threshold. On a sample drawn from one month, that threshold needed 43 "
  "components. On a sample drawn from five months spread across the full "
  "date range, it needed 59. The gap makes sense once you think about what "
  "the embedding encodes: a wider span of months carries more distinct job "
  "names and, presumably, more distinct applications, so capturing the same "
  "share of variance takes more components. The number is not fixed until a "
  "fit is done against the full 38-month dataset; treat 59 as current, not "
  "final.")
fig("nb02_cell22_img3.png",
    "Figure 3.1. Cumulative explained variance of F-DATA's job-name embedding "
    "against number of PCA components retained, computed on a 5-file sample. "
    "The dashed line marks the 90% threshold used to select the component count.")
p("I also decided, after an earlier round of discussion, to use the "
  "PCA-reduced embedding everywhere — for the tree-based models and for the "
  "neural network architectures alike — rather than giving the deep learning "
  "models the option of the raw 384-dimensional vector, which the original "
  "proposal allowed for. The reason is practical rather than statistical: "
  "loading the raw embedding for the full F-DATA dataset measures at over 100 "
  "gigabytes of resident memory, against roughly 27 gigabytes for every other "
  "column combined, because each row's embedding is stored as an individually "
  "allocated array rather than one contiguous block. Section 5.2 covers this "
  "measurement and the loading strategy it led to in more detail.")

h2("3.5 Feature Vetting: PM100's Missing-by-Design Fields")
p("Three PM100 columns are Tier A candidates and all three have missing "
  "values, but at very different rates: num_tasks is missing in 4.1 percent "
  "of jobs, req_nodes in 88.6 percent, and threads_per_core in 99.6 percent. "
  "Treating all three the same way — impute and move on — would have missed "
  "what turned out to be a fairly consequential distinction, so I worked "
  "through a four-part check on each candidate before deciding whether its "
  "derived feature belonged in the active Tier A set.")
p("The first check is simply whether the missingness reflects an optional "
  "field or an incomplete one. req_nodes, when populated, holds a specific "
  "node identifier — values like [763] or [878] — which is Slurm's explicit "
  "node-pinning flag (--nodelist). Most jobs do not ask for a particular "
  "machine, so the field is null by design for them, not because logging "
  "failed. threads_per_core is the same story: its only non-null values are "
  "1, 2, or 4, matching Slurm's explicit hyperthreading override "
  "(--threads-per-core), a setting almost nobody touches. num_tasks has no "
  "such documented optional-flag semantics; its gaps look like ordinary "
  "incomplete logging.")
p("The second check is effect size against the three prediction targets, "
  "using median comparisons rather than leaning on p-values alone — at "
  "roughly 180,000 completed jobs, a Mann-Whitney test will report "
  "significance for almost any difference, trivial or not. All three derived "
  "flags (node_pinned, threads_per_core_set, and num_tasks_missing) showed "
  "real median differences. threads_per_core_set showed the largest of the "
  "three by a wide margin: median run_time roughly four times higher, median "
  "memory roughly four times lower, when the flag is set.")
fig("nb02_cell11_img1.png",
    "Figure 3.2. Median run_time, mem_alloc, and power for each candidate flag, "
    "true versus false, on a log scale. threads_per_core_set shows the largest "
    "separation of the three despite being the rarest.")
p("The third check is redundancy: does an existing categorical column already "
  "carry the same information as the flag being considered? I cross-tabulated "
  "threads_per_core_set against partition, quality-of-service class, and GPU "
  "allocation. None of them explained it — the flag is set almost exclusively "
  "within the system's single dominant partition and QoS class, which is "
  "where nearly every job lives regardless, and GPU allocation among "
  "flagged jobs was close to the overall rate. Ruling out redundancy raised "
  "my confidence in the flag rather than lowering it, which is why the fourth "
  "check mattered more than I expected going in.")
p("The fourth check is user concentration: is the signal broad, or is it a "
  "handful of individual users' behavior showing up as if it were a general "
  "pattern? Both datasets are naturally dominated by a small number of heavy "
  "users — PM100's top five account for 45.4 percent of all completed jobs, "
  "F-DATA's top five for 41.4 percent in a five-month sample — so some "
  "concentration is expected and not automatically disqualifying. The "
  "question is how much concentration is too much. node_pinned turned out to "
  "be far more concentrated than that baseline: only 19 users ever set it, "
  "and the top five account for 99.3 percent of every pinned job, 73.5 "
  "percent from one user alone. threads_per_core_set was similarly narrow: "
  "only 6 users, 85 percent of occurrences from a single one of them. "
  "num_tasks_missing, by contrast, spans 101 distinct users with the top "
  "five accounting for 47.1 percent — close enough to the dataset's own "
  "baseline that it reads as a genuine pattern rather than a few users' "
  "signature.")
fig("nb02_cell17_img2.png",
    "Figure 3.3. Top-1 and top-5 user share of jobs for each candidate flag, "
    "compared against the dataset's overall baseline concentration (dashed "
    "line). node_pinned and threads_per_core_set both sit well above the "
    "baseline; num_tasks_missing sits close to it.")
p("Only num_tasks_missing passes all four checks, and it is the feature "
  "promoted into the active Tier A set. node_pinned and threads_per_core_set "
  "are both computed and kept in the processed dataset — along with the raw "
  "req_nodes and threads_per_core columns — but excluded from the general "
  "feature set, reserved for a narrower, explicitly scoped analysis later "
  "(a power-specific submodel is the most likely candidate, given "
  "threads_per_core_set's relationship to SMT settings) where the "
  "single-user-concentration caveat can be addressed directly rather than "
  "quietly inherited by every model that uses Tier A. threads_per_core_set "
  "having the largest raw effect size of the three, and still being the one "
  "excluded, is worth stating plainly: effect size on its own was not "
  "sufficient evidence, and would have been misleading without the "
  "concentration check that followed it.")

p("That decision covers three candidate features. It says nothing about "
  "whether the same problem was already sitting undetected somewhere in the "
  "feature set built before this vetting process existed, so I went back "
  "and ran the same concentration check against every PM100 Tier A column "
  "that predates it. One failed outright: req_switch, a Slurm switch-count "
  "field already in active use, has only 13 non-zero values across 180,310 "
  "completed jobs, and all 13 come from a single user — a worse "
  "concentration ratio than threads_per_core_set, which had already been "
  "excluded for the same reason. req_switch has now been moved to the same "
  "reserved category, with a derived req_switch_set flag preserved "
  "alongside it for consistency. The rest of the audit came back clean: "
  "shared's two populated categories each hold tens of thousands of jobs, "
  "nowhere near single-user territory, and the smallest categories of qos "
  "and partition, while occasionally tiny, are legitimate multi-valued "
  "scheduling fields rather than purpose-built rare flags, so ordinary "
  "small-category noise there is not the same concern. F-DATA's one "
  "comparable column, jobenv_req, showed a milder version of the same "
  "pattern — its minority value spans 21 users with a 39.2 percent top-user "
  "share, above F-DATA's own baseline concentration of 16.0 percent but "
  "well short of PM100's excluded flags — and is left active for now, "
  "flagged for a second look once more months of F-DATA are loaded.")

p("There is a second, broader question this whole exercise raised that is "
  "worth stating on its own rather than folding into the feature-by-feature "
  "discussion above: given that a handful of users account for close to "
  "half of all jobs in both datasets, does the chronological train/test "
  "split planned for later chapters risk giving an unrepresentative slice "
  "of any one heavy user's jobs to training versus testing? I checked this "
  "for PM100's five heaviest users directly, rather than leaving it as an "
  "open question. All five are active across 97 to 100 percent of the "
  "dataset's full 159-day span, with one exception reaching 76 percent "
  "after starting roughly a month into the window — none of them are "
  "concentrated in a narrow slice of time that a chronological split could "
  "isolate entirely on one side. That is a reassuring result, but it is "
  "specific to PM100. F-DATA spans 38 months rather than 6, which is a "
  "much longer window for a user's activity to plausibly shift, taper off, "
  "or start partway through before the dataset ends, and that check has "
  "not yet been run there. The Feature Vetting method described in this "
  "section should be treated as an ongoing discipline for this thesis "
  "rather than a step that was completed once and closed — the fact that "
  "it caught req_switch on a second pass, after already being applied "
  "once, is itself the argument for continuing to apply it as new features "
  "and the full-scale dataset come into the pipeline.")

h2("3.6 Handling Incomplete Jobs")
p("Both datasets record jobs that failed, were cancelled, or hit resource "
  "limits before completing. Their duration and power figures reflect the "
  "failure, not ordinary workload behavior, so both datasets are filtered to "
  "completed jobs only before any target is computed. F-DATA's exit-state "
  "field is a clean binary and excludes 5.1–5.6 percent of jobs depending on "
  "the sample; PM100's job-state field has six values — completed, failed, "
  "cancelled, timeout, out-of-memory, and node failure — and excludes a "
  "notably larger 21.0 percent. That gap between the two systems' failure "
  "rates is itself worth stating explicitly rather than treating the two "
  "datasets as symmetric on this point.")

# ============================================================
doc.add_page_break()
h1("Chapter 5 (partial): Experimental Environment and Reproducibility")
p("This chapter's final form will also cover model training details once "
  "those exist. The material below documents the computing environment and "
  "the memory constraints discovered while building the data pipeline, which "
  "belong here regardless of what else the chapter ends up containing.")

h2("5.1 Computing Environment")
p("All experiments in this thesis, at least through the feature-engineering "
  "stage, run on a single workstation: an Intel i9-10900X (10 cores, 20 "
  "threads), 125 gigabytes of RAM, and an NVIDIA RTX 3090 with 24 gigabytes "
  "of VRAM. The environment is managed with uv rather than plain pip, "
  "pinned to Python 3.12 rather than the newer 3.13 available on the "
  "system. That choice was not arbitrary: numba, a dependency of the SHAP "
  "interpretability library used later in this work, does not yet build "
  "cleanly under Python 3.13. More specifically, resolving the full "
  "dependency set under the newer Python version led the package resolver "
  "to select numba 0.53.1, a version old enough that its own dependency, "
  "llvmlite, refuses to build at all above Python 3.9. The fix needed an "
  "explicit lower bound on numba in the project's dependency list; without "
  "it, the resolver's preference for a version whose metadata claims "
  "unbounded compatibility, even when that claim is stale, wins over a "
  "modern release that correctly declares its actual constraints. PyTorch "
  "is installed from the CUDA 13.0 build, matching the workstation's driver "
  "version exactly, and its GPU detection was confirmed directly rather than "
  "assumed from the installation succeeding.")

h2("5.2 Memory Constraints at Full Scale")
p("F-DATA's 38 monthly files total roughly 26 gigabytes on disk, comfortably "
  "within the workstation's storage, but disk size is not the same question "
  "as memory footprint once loaded. Measuring one month's file directly: "
  "loading all 45 columns costs 5.54 gigabytes of resident memory, while "
  "loading the same file without the embedding column costs 1.39 gigabytes. "
  "Scaled to the full 25,866,900-row dataset, that difference becomes "
  "significant — upwards of 100 gigabytes with the embedding included, "
  "against roughly 27 gigabytes without it. Given 125 gigabytes of total "
  "RAM, the first scenario would leave little headroom for anything else "
  "running on the machine, let alone for the model training this dataset "
  "exists to support.")
p("The inefficiency is specific to how the embedding is stored: each row's "
  "384-dimensional vector is an individually allocated NumPy array rather "
  "than a slice of one contiguous block, which carries far more overhead per "
  "row than the memory the numbers themselves would need. The fix separates "
  "fitting the PCA transform from applying it. The transform is fit once, on "
  "a sample small enough to hold comfortably in memory, and the fitted model "
  "is then applied to one monthly file at a time in a loop, so the raw "
  "embedding for more than one month's worth of rows is never resident at "
  "once. A function that loads every other column across all 38 files "
  "without touching the embedding at all covers the case where "
  "embedding-derived features are not needed for a given step.")

h2("5.3 Notebook Execution and Monitoring")
p("Early runs of the data-loading notebook appeared to hang for extended "
  "periods with no visible output, which turned out to have two distinct "
  "causes rather than one. The first was calling a generic describe() "
  "across a dataframe that still included the embedding column: pandas "
  "attempts to compute uniqueness and mode statistics on that column, which "
  "means comparing unhashable NumPy arrays across a million rows, and the "
  "fallback path for that is slow rather than an outright error. The second "
  "was treating PM100's node_power_consumption as if it were a single "
  "number per job, when it is in fact a time series — power sampled roughly "
  "every 20 seconds for the duration of the job, which is exactly the kind "
  "of sequence data the LSTM and TCN architectures in Chapter 4 are meant to "
  "use. Both were fixed at the source rather than worked around: the "
  "embedding column is excluded from any generic summary call, and any "
  "scalar treatment of the power column reduces the sequence to a per-job "
  "mean explicitly first.")
p("Notebooks are now executed with papermill rather than plain jupyter "
  "nbconvert, specifically because nbconvert only writes output once "
  "execution finishes entirely, which gives no way to distinguish a slow "
  "cell from a stalled one while it is running. papermill logs each cell's "
  "start and end as it happens, which is what actually resolved the "
  "hanging-versus-slow question above rather than guesswork based on "
  "process CPU time.")

doc.save(OUT_PATH)
print("saved to", OUT_PATH)
